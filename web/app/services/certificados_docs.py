# web/app/services/certificados_docs.py
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document as DocxDocument
from flask import current_app

from ..models import CertificadoLaboral
from .docx_engine import replace_placeholders


def _fmt_fecha_es(d: date) -> str:
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _only_digits(s: str) -> str:
    return re.sub(r"\D+", "", s or "")


def _split_rut_dv(rut: str | None, dv: str | None) -> tuple[str, str]:
    """
    Devuelve (rut_digits, dv_clean) robusto.
    Soporta:
      - rut="76.772.137", dv="4"
      - rut="76.772.137-4", dv=None
      - rut="767721374", dv=None        (último dígito asumido como DV)
      - rut="76772137K", dv=None        (último char K asumido como DV)
    """
    rut_raw = (rut or "").strip()
    dv_clean = (dv or "").strip().upper()

    if not rut_raw:
        return ("", dv_clean)

    # Caso 1: viene con guion (lo más confiable)
    if "-" in rut_raw:
        left, right = rut_raw.split("-", 1)
        rut_digits = _only_digits(left)
        dv_guess = (right or "").strip().upper()
        if not dv_clean and dv_guess:
            dv_clean = dv_guess[:1]  # 1 char
        return (rut_digits, dv_clean)

    # Caso 2: si no hay dv explícito, intentar extraer último char si parece DV
    if not dv_clean:
        # Normaliza quitando puntos/espacios
        compact = re.sub(r"[.\s]", "", rut_raw).strip()
        if len(compact) >= 2:
            last = compact[-1].upper()
            head = compact[:-1]
            # DV puede ser 0-9 o K
            if re.fullmatch(r"[0-9K]", last) and _only_digits(head):
                dv_clean = last
                rut_digits = _only_digits(head)
                return (rut_digits, dv_clean)

    # Caso 3: fallback: todo como rut (sin dv)
    rut_digits = _only_digits(rut_raw)
    return (rut_digits, dv_clean)


def _format_rut_with_dots(rut: str | None, dv: str | None) -> str:
    """
    Formatea RUT chileno:
      16232010 + 6  => 16.232.010-6
      "76.772.137-4" => 76.772.137-4
      "767721374" (rut+dv) => 76.772.137-4
    """
    rut_digits, dv_clean = _split_rut_dv(rut, dv)

    if not rut_digits:
        return ""

    # Insertar puntos cada 3 desde la derecha
    parts = []
    r = rut_digits
    while len(r) > 3:
        parts.append(r[-3:])
        r = r[:-3]
    parts.append(r)
    rut_fmt = ".".join(reversed(parts))

    return f"{rut_fmt}-{dv_clean}" if dv_clean else rut_fmt



_SPACE_RE = re.compile(r"[ \t]{2,}")
_SPACE_BEFORE_PUNCT_RE = re.compile(r"\s+([,.;:])")
_SPACE_BEFORE_CLOSE_RE = re.compile(r"\s+(\))")


def _normalize_text_spacing(s: str) -> str:
    """
    - Colapsa dobles espacios
    - Quita espacios antes de puntuación: "Louis ," => "Louis,"
    - Quita espacios antes de ")"
    """
    if not s:
        return s
    s = _SPACE_RE.sub(" ", s)
    s = _SPACE_BEFORE_PUNCT_RE.sub(r"\1", s)
    s = _SPACE_BEFORE_CLOSE_RE.sub(r"\1", s)
    return s


def _cleanup_doc_spacing(doc: DocxDocument) -> None:
    """
    Limpieza suave post-reemplazo:
    recorre runs para no reventar el formato del párrafo completo.
    """
    def _clean_runs(paragraph) -> None:
        for run in paragraph.runs:
            if run.text:
                run.text = _normalize_text_spacing(run.text)

    # Párrafos normales
    for p in doc.paragraphs:
        _clean_runs(p)

    # Tablas (si la plantilla tiene)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _clean_runs(p)


def build_certif_antiguedad_docx(
    cert: CertificadoLaboral,
    output_path: Path,
    overrides: dict | None = None,
) -> None:
    """
    Genera Certificado de Antigüedad DOCX desde plantilla corporativa.
    Plantilla: web/app/document_templates/certificados/certif_antiguedad.docx
    """
    overrides = overrides or {}

    tpl = Path(current_app.root_path) / "document_templates" / "certificados" / "certif_antiguedad.docx"
    if not tpl.exists():
        raise FileNotFoundError(f"No se encontró plantilla: {tpl}")

    t = cert.trabajador
    c = cert.contrato

    obra = (
        getattr(cert, "obra", None)
        or (getattr(c, "obra", None) if c else None)
        or (getattr(t, "obra", None) if t else None)
    )

    empleador = (
        getattr(cert, "empleador", None)
        or (getattr(c, "empleador", None) if c else None)
        or (getattr(t, "empleador_preferente", None) if t else None)
    )

    if not t:
        raise RuntimeError("Certificado sin trabajador asociado.")
    if not empleador:
        raise RuntimeError("No se pudo resolver empleador para el certificado (cert/contrato/trabajador).")

    # Cargo (con fallback conservador)
    cargo_nombre = "—"
    try:
        if c and getattr(c, "cargo", None):
            cargo_nombre = c.cargo.nombre or "—"
        elif getattr(t, "cargo", None):
            cargo_nombre = t.cargo.nombre or "—"
    except Exception:
        pass

    # Fecha ingreso (con fallback conservador)
    if c and getattr(c, "fecha_inicio", None):
        fecha_ingreso = c.fecha_inicio
    elif getattr(t, "fecha_ingreso_empresa", None):
        fecha_ingreso = t.fecha_ingreso_empresa
    else:
        fecha_ingreso = cert.fecha_emision or date.today()

    fecha_emision = cert.fecha_emision or date.today()

    # Tipo de duración del contrato (ya correcto en femenino: "duración indefinida")
    tipo_duracion = "indefinida"

    if c and (c.tipo_contrato or "").strip():
        tt = (c.tipo_contrato or "").lower()

        if "plazo" in tt or "fijo" in tt:
            if c.fecha_termino:
                tipo_duracion = f"a plazo fijo, hasta el {_fmt_fecha_es(c.fecha_termino)}"
            else:
                tipo_duracion = "a plazo fijo"
        elif "obra" in tt or "faena" in tt:
            tipo_duracion = "por obra o faena"
        elif "indef" in tt:
            tipo_duracion = "indefinida"

    # RUTs con formato chileno
    rut_trabajador = _format_rut_with_dots(getattr(t, "rut", None), getattr(t, "dv", None))
    rut_empleador = _format_rut_with_dots(getattr(empleador, "rut", None), getattr(empleador, "dv", None))

    # Apellidos/nombres (si ap_materno es None => "", luego limpiamos espacios/puntuación)
    nombres = (t.nombres or "").strip()
    ap_paterno = (t.ap_paterno or "").strip()
    ap_materno_raw = (t.ap_materno or "").strip()

    ap_materno = f" {ap_materno_raw}" if ap_materno_raw else ""

    mapping = {
        # Empleador
        "EMPLEADOR_RAZON_SOCIAL": (getattr(empleador, "razon_social", "") or ""),
        "EMPLEADOR_RUT": rut_empleador or (getattr(empleador, "rut", "") or ""),

        # Trabajador
        "TRABAJADOR_NOMBRES": nombres,
        "TRABAJADOR_AP_PATERNO": ap_paterno,
        "TRABAJADOR_AP_MATERNO": ap_materno,
        "TRABAJADOR_RUT": rut_trabajador or (getattr(t, "rut_completo", None) or ""),

        # Laboral
        "CARGO_NOMBRE": cargo_nombre,
        "CONTRATO_FECHA_INGRESO": _fmt_fecha_es(fecha_ingreso),

        # Obra
        "OBRA_NOMBRE": (getattr(obra, "nombre", None) or "—"),
        "OBRA_DIRECCION": (getattr(obra, "direccion", None) or "—"),
        "OBRA_COMUNA": (getattr(obra, "comuna", None) or "—"),

        # Certificado
        "CERTIFICADO_OBJETIVO": (cert.objetivo or "los fines que estime convenientes"),
        "CERTIFICADO_FECHA": _fmt_fecha_es(fecha_emision),

        # Tipo duración
        "TIPO_DURACION": tipo_duracion,
    }

    # Overrides pisan mapping (igual que desvinculaciones)
    for k, v in overrides.items():
        mapping[str(k)] = "" if v is None else str(v)

    doc = DocxDocument(str(tpl))
    replace_placeholders(doc, mapping)

    # ✅ Fix: limpia dobles espacios + "espacio antes de coma" en todo el doc
    _cleanup_doc_spacing(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
