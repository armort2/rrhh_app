# web/app/services/desvinculaciones_docs.py
from __future__ import annotations

from datetime import date
from pathlib import Path
import calendar
import re

from flask import current_app

import sqlalchemy as sa
from sqlalchemy import or_

from docx import Document  # python-docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ..models import (
    Desvinculacion,
    Trabajador,
    Contrato,
    Inasistencia,
    DesvinculacionDocInputs,  # ✅ FALTABA (usado en upsert/get)
)
from ..utils import fecha_larga_es
from ..services.docx_engine import replace_placeholders
from ..services.pdf_convert import convert_docx_to_pdf
from ..extensions import db


_MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
    7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}
_DIAS_SEMANA_ES = {
    0: "lunes", 1: "martes", 2: "miércoles", 3: "jueves", 4: "viernes", 5: "sábado", 6: "domingo",
}


def _fmt(d: date | None) -> str:
    return d.strftime("%d-%m-%Y") if d else ""


def _join_es(items: list[str]) -> str:
    items = [x for x in items if x]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} y {items[1]}"
    return ", ".join(items[:-1]) + " y " + items[-1]


def _month_bounds(d: date) -> tuple[date, date]:
    last_day = calendar.monthrange(d.year, d.month)[1]
    return date(d.year, d.month, 1), date(d.year, d.month, last_day)


def _consecutive_runs(dates_sorted: list[date]) -> list[list[date]]:
    runs: list[list[date]] = []
    run: list[date] = []
    for dt in dates_sorted:
        if not run:
            run = [dt]
            continue
        prev = run[-1]
        if (dt - prev).days == 1:
            run.append(dt)
        else:
            runs.append(run)
            run = [dt]
    if run:
        runs.append(run)
    return runs


def _templates_root() -> Path:
    base_dir = Path(current_app.root_path)  # /app/app
    rel = current_app.config.get("DESVINCULACION_TEMPLATES_DIR", "document_templates/desvinculacion")
    return (base_dir / rel).resolve()


def _resolve_template(template_name_or_relpath: str) -> Path:
    if not template_name_or_relpath:
        raise RuntimeError("Nombre de plantilla vacío.")

    root = _templates_root()
    candidate = (root / template_name_or_relpath).resolve()

    if root not in candidate.parents and candidate != root:
        raise RuntimeError("Ruta de plantilla inválida (fuera del directorio permitido).")

    if not candidate.exists():
        raise RuntimeError(f"No existe la plantilla DOCX: {candidate}")

    return candidate


def _to_str_dict(d: dict) -> dict[str, str]:
    out: dict[str, str] = {}
    for k, v in (d or {}).items():
        out[str(k)] = "" if v is None else str(v)
    return out


# --- helpers tabla finiquito dinámica ---

_MONEY_ONLY_DIGITS_RE = re.compile(r"[^\d]")


def _money_to_int(val: str | None) -> int:
    """Convierte '1.250.000' / '1250000' / '' -> int (pesos)."""
    s = (val or "").strip()
    if not s:
        return 0
    digits = _MONEY_ONLY_DIGITS_RE.sub("", s)
    return int(digits) if digits else 0


def _has_amount(val: str | None) -> bool:
    return _money_to_int(val) > 0


# ============================================================
# ✅ NUEVO: Bloques condicionales para carta (Art. 161)
# ============================================================

def _build_bloque_aviso_previo(ctx: dict[str, str]) -> str:
    """
    Devuelve el párrafo (o vacío) según FINIQUITO_AVISO_PREVIO.
    Nota: el monto ya viene formateado como '1.250.000' por tu flujo.
    """
    monto = (ctx.get("FINIQUITO_AVISO_PREVIO") or "").strip()
    if not _has_amount(monto):
        return ""

    # Texto corporativo, conservador y estándar.
    return (
        "En consecuencia, se le pagará la indemnización sustitutiva del aviso previo, "
        f"por la suma de $ {monto}.-, la que será incluida en el finiquito respectivo."
    )


def _build_bloque_anios_servicio(ctx: dict[str, str]) -> str:
    """
    Devuelve el párrafo (o vacío) según FINIQUITO_ANIOS_SERVICIOS.
    """
    monto = (ctx.get("FINIQUITO_ANIOS_SERVICIOS") or "").strip()
    if not _has_amount(monto):
        return ""

    return (
        "Asimismo, se le pagará la indemnización por años de servicio, "
        f"por la suma de $ {monto}.-, la que será incluida en el finiquito respectivo."
    )


def _apply_bloques_carta(ctx: dict[str, str]) -> dict[str, str]:
    """
    Inyecta BLOQUE_AVISO_PREVIO y BLOQUE_ANIOS_SERVICIO al contexto final.
    Si no hay monto => deja vacío => el Word queda sin ese párrafo.
    """
    ctx["BLOQUE_AVISO_PREVIO"] = _build_bloque_aviso_previo(ctx)
    ctx["BLOQUE_ANIOS_SERVICIO"] = _build_bloque_anios_servicio(ctx)
    return ctx


def _iter_paragraphs(doc: Document):
    """Itera párrafos del doc incluyendo los que están dentro de tablas."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _find_paragraph_with_marker(doc: Document, marker: str):
    for p in _iter_paragraphs(doc):
        if marker in (p.text or ""):
            return p
    return None


def _clear_paragraph_marker(paragraph, marker: str) -> None:
    if paragraph.runs:
        for r in paragraph.runs:
            r.text = r.text.replace(marker, "")
    else:
        paragraph.text = (paragraph.text or "").replace(marker, "")


def _delete_paragraph(paragraph) -> None:
    """
    Elimina un párrafo del DOCX (body o tabla).
    python-docx no tiene API pública para esto, así que usamos el XML.
    """
    try:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None  # type: ignore[attr-defined]
    except Exception:
        # Si algo falla, preferimos no romper la generación completa.
        pass


def _prune_optional_paragraphs(doc: Document, mapping: dict[str, str]) -> None:
    """
    Borra párrafos completos que contienen marcadores opcionales
    cuando el valor a inyectar viene vacío.

    Importante: esto se ejecuta ANTES de replace_placeholders(),
    para encontrar el marcador intacto.
    """
    optional_keys = ["BLOQUE_AVISO_PREVIO", "BLOQUE_ANIOS_SERVICIO"]

    for key in optional_keys:
        marker = f"{{{{{key}}}}}"
        value = (mapping.get(key) or "").strip()

        if value:
            continue  # aplica, se queda

        # Recorremos una lista "materializada" para poder borrar sin romper el iterador.
        for p in list(_iter_paragraphs(doc)):
            if marker in (p.text or ""):
                # Seguridad: si el párrafo tiene más texto además del marcador,
                # en vez de borrar, solo limpiamos el marcador.
                only_marker = (p.text or "").strip() == marker
                if only_marker:
                    _delete_paragraph(p)
                else:
                    _clear_paragraph_marker(p, marker)


def _set_cell_vertical_center(cell) -> None:
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        pass


def _set_paragraph_style(p, align: WD_ALIGN_PARAGRAPH, bold: bool = False, size_pt: int | None = None) -> None:
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)

    # Aplicar formato a todos los runs existentes; si no hay, crear uno.
    if not p.runs:
        p.add_run("")

    for r in p.runs:
        if bold:
            r.bold = True
        if size_pt is not None:
            r.font.size = Pt(size_pt)


def _set_cell_text(
    cell,
    text: str,
    *,
    align: WD_ALIGN_PARAGRAPH,
    valign_center: bool = True,
    bold: bool = False,
    size_pt: int | None = None,
) -> None:
    # limpiar contenido previo
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

    # asegurar un solo párrafo "controlado"
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.text = text or ""

    if valign_center:
        _set_cell_vertical_center(cell)

    _set_paragraph_style(p, align=align, bold=bold, size_pt=size_pt)


def _insert_table_after_paragraph(doc: Document, paragraph, rows: list[tuple[str, str, str]]):
    """
    Inserta una tabla (3 columnas) justo después del párrafo recibido.
    Compatible con:
      - Document body (add_table sin width)
      - Celdas / BlockItemContainer (add_table exige width)
    """
    container = paragraph._parent  # Document body o celda

    # --- ancho útil real ---
    try:
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
    except Exception:
        usable_width = Inches(6.0)

    # ✅ Crear tabla con fallback según el tipo de container
    try:
        table = container.add_table(rows=0, cols=3)  # type: ignore[call-arg]
    except TypeError:
        # BlockItemContainer (celdas) suele exigir width
        table = container.add_table(rows=0, cols=3, width=usable_width)  # type: ignore[call-arg]

    table.style = "Table Grid"

    # Fijar anchos (evita que LibreOffice “autofitee” distinto en PDF)
    try:
        table.autofit = False
    except Exception:
        pass

    # Proporciones
    concept_w = int(usable_width * 0.72)
    currency_w = int(usable_width * 0.06)
    amount_w = usable_width - concept_w - currency_w

    table.columns[0].width = concept_w
    table.columns[1].width = currency_w
    table.columns[2].width = amount_w

    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for concepto, moneda, monto in rows:
        concepto_norm = (concepto or "").strip()
        concepto_key = concepto_norm.upper()

        is_total_row = (concepto_key == "TOTAL A PAGO")
        is_subtotal_row = (concepto_key == "SUB TOTAL")
        is_total_desc_row = (concepto_key == "TOTAL DESCUENTOS")
        is_mid_total_row = is_subtotal_row or is_total_desc_row

        cells = table.add_row().cells

        # Re-asegurar anchos por celda (LO lo respeta mejor)
        try:
            cells[0].width = concept_w
            cells[1].width = currency_w
            cells[2].width = amount_w
        except Exception:
            pass

        # 1) Concepto
        if is_total_row:
            _set_cell_text(cells[0], concepto_norm, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size_pt=14)
        elif is_mid_total_row:
            _set_cell_text(cells[0], concepto_norm, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size_pt=12)
        else:
            _set_cell_text(cells[0], concepto_norm, align=WD_ALIGN_PARAGRAPH.LEFT)

        # 2) $
        _set_cell_text(
            cells[1], moneda, align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True if is_total_row else False,
            size_pt=14 if is_total_row else None
        )

        # 3) Monto
        if is_total_row:
            _set_cell_text(cells[2], monto, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size_pt=14)
        elif is_mid_total_row:
            _set_cell_text(cells[2], monto, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size_pt=12)
        else:
            _set_cell_text(cells[2], monto, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # mover tabla inmediatamente después del párrafo
    paragraph._p.addnext(table._tbl)  # type: ignore[attr-defined]


def _should_include_subtotal(sumandos_con_monto: int, has_descuentos: bool) -> bool:
    """
    ✅ Regla nueva:
    - Si NO hay descuentos -> NO mostramos Sub Total (evita duplicar TOTAL A PAGO).
    - Si hay descuentos y hay al menos 2 ítems con monto -> sí mostramos Sub Total.
    """
    return has_descuentos and (sumandos_con_monto >= 2)


def _build_finiquito_valores_rows(ctx: dict[str, str]) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []

    # --- SUMANDOS ---
    sumandos: list[tuple[str, str]] = []

    dias = (ctx.get("FINIQUITO_DIAS_FERIADO_A_PAGO") or "").strip()
    feriado_monto = ctx.get("FINIQUITO_FERIADO_PROPORCIONAL", "") or ""

    concepto_feriado = "Indemnización Feriado Proporcional"
    if dias:
        concepto_feriado += f" ({dias} días hábiles)"
    sumandos.append((concepto_feriado, feriado_monto))

    if _has_amount(ctx.get("FINIQUITO_AVISO_PREVIO")):
        sumandos.append(("Indemnización falta de Aviso Previo", ctx.get("FINIQUITO_AVISO_PREVIO", "")))

    if _has_amount(ctx.get("FINIQUITO_ANIOS_SERVICIOS")):
        sumandos.append(("Indemnización por Años de Servicios", ctx.get("FINIQUITO_ANIOS_SERVICIOS", "")))

    if _has_amount(ctx.get("FINIQUITO_INDEM_VOLUNTARIA")):
        sumandos.append(("Indemnización Voluntaria", ctx.get("FINIQUITO_INDEM_VOLUNTARIA", "")))

    otros_label = (ctx.get("FINIQUITO_OTROS") or "").strip()
    otros_monto = ctx.get("FINIQUITO_MONTO_OTROS") or ""
    if otros_label and _has_amount(otros_monto):
        sumandos.append((otros_label, otros_monto))

    # --- DESCUENTOS (primero armamos la lista para decidir Sub Total) ---
    descuentos: list[tuple[str, str]] = []

    if _has_amount(ctx.get("FINIQUITO_APORTE_EMP_SEG_CES")):
        descuentos.append(("(MENOS) Aporte empleador al Seg. Cesantía", ctx.get("FINIQUITO_APORTE_EMP_SEG_CES", "")))

    d1_label = (ctx.get("FINIQUITO_DESCTO_1") or "").strip()
    d1_monto = (ctx.get("FINIQUITO_MONTO_DESCTO_1") or "").strip()
    if d1_label and _has_amount(d1_monto):
        descuentos.append((f"(MENOS) {d1_label}", d1_monto))

    d2_label = (ctx.get("FINIQUITO_DESCTO_2") or "").strip()
    d2_monto = (ctx.get("FINIQUITO_MONTO_DESCTO_2") or "").strip()
    if d2_label and _has_amount(d2_monto):
        descuentos.append((f"(MENOS) {d2_label}", d2_monto))

    has_descuentos = len(descuentos) > 0

    # 1) Imprimir sumandos
    for concepto, monto in sumandos:
        rows.append((concepto, "$", f"{monto}.-" if monto else ""))

    # 2) ✅ Sub Total solo si hay descuentos (evita duplicar TOTAL A PAGO cuando es “solo sumas”)
    sumandos_con_monto = sum(1 for _, monto in sumandos if _has_amount(monto))
    if _should_include_subtotal(sumandos_con_monto, has_descuentos):
        sub = (ctx.get("FINIQUITO_SUBTOTAL") or "").strip()
        rows.append(("Sub Total", "$", f"{sub}.-" if sub else ""))

    # 3) Imprimir descuentos y total descuentos
    for concepto, monto in descuentos:
        rows.append((concepto, "$", f"{monto}.-" if monto else ""))

    if descuentos:
        td = (ctx.get("FINIQUITO_TOTAL_DESCTOS") or "").strip()
        rows.append(("Total Descuentos", "$", f"{td}.-" if td else ""))

    # 4) Total a pago
    total = (ctx.get("FINIQUITO_TOTAL_PAGO") or "").strip()
    rows.append(("TOTAL A PAGO", "$", f"{total}.-" if total else ""))

    return rows


def _inject_finiquito_table(doc: Document, ctx: dict[str, str]) -> None:
    marker = "{{FINIQUITO_VALORES}}"
    p = _find_paragraph_with_marker(doc, marker)
    if not p:
        return

    _clear_paragraph_marker(p, marker)
    rows = _build_finiquito_valores_rows(ctx)
    _insert_table_after_paragraph(doc, p, rows)


def _render_docx_template(template_path: Path, mapping: dict[str, str], output_path: Path) -> Path:
    doc = Document(str(template_path))

    # ✅ NUEVO: borrar párrafos opcionales si el bloque viene vacío
    _prune_optional_paragraphs(doc, mapping)

    replace_placeholders(doc, mapping)
    _inject_finiquito_table(doc, mapping)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _fmt_dia_mes_anio_es(d: date, include_weekday: bool = False) -> str:
    mes = _MESES_ES.get(d.month, str(d.month))
    base = f"{d.day} de {mes} de {d.year}"
    if include_weekday:
        wd = _DIAS_SEMANA_ES.get(d.weekday(), "")
        return f"{wd} {base}".strip()
    return base


def _fmt_larga(d: date | None) -> str:
    if not d:
        return ""
    return _fmt_dia_mes_anio_es(d, include_weekday=False)


def _format_rut(rut_digits: str, dv: str) -> str:
    rut_digits = re.sub(r"\D", "", rut_digits or "")
    dv = (dv or "").strip().upper()
    if not rut_digits or not dv:
        return ""
    body = f"{int(rut_digits):,}".replace(",", ".")
    return f"{body}-{dv}"


def _base_context(desv: Desvinculacion) -> dict:
    t: Trabajador = desv.trabajador
    c: Contrato = desv.contrato
    causal = desv.causal

    obra = getattr(c, "obra", None)
    emp = getattr(c, "empleador", None)
    cargo = getattr(c, "cargo", None)

    trabajador_nombres = (t.nombres or "").strip()
    trabajador_rut = _format_rut(getattr(t, "rut", "") or "", getattr(t, "dv", "") or "")

    afp_nombre = ""
    try:
        afp_nombre = (t.afp.nombre or "") if getattr(t, "afp", None) else ""
    except Exception:
        afp_nombre = ""

    salud_nombre = ""
    try:
        salud_nombre = (t.salud.nombre or "") if getattr(t, "salud", None) else ""
    except Exception:
        salud_nombre = ""

    empleador_razon = (getattr(emp, "razon_social", "") or "") if emp else ""
    empleador_rut = (getattr(emp, "rut", "") or "") if emp else ""
    empleador_dir = (getattr(emp, "direccion", "") or "") if emp else ""
    empleador_comuna = (getattr(emp, "comuna", "") or "") if emp else ""

    rep_legal_nombre = (getattr(emp, "nombre_rep_legal", "") or "") if emp else ""
    rep_legal_rut = (getattr(emp, "rut_rep_legal", "") or "") if emp else ""

    obra_nombre = (getattr(obra, "nombre", "") or "") if obra else ""
    obra_dir = (getattr(obra, "direccion", "") or "") if obra else ""
    obra_comuna = (getattr(obra, "comuna", "") or "") if obra else ""

    cargo_nombre = (getattr(cargo, "nombre", "") or "") if cargo else ""

    fecha_emision = date.today()

    ctx = {
        "TRABAJADOR_NOMBRES": trabajador_nombres,
        "TRABAJADOR_AP_PATERNO": (t.ap_paterno or ""),
        "TRABAJADOR_AP_MATERNO": (t.ap_materno or ""),
        "TRABAJADOR_RUT": trabajador_rut,
        "TRABAJADOR_DIRECCION": (t.direccion or ""),
        "TRABAJADOR_COMUNA": (t.comuna or ""),
        "TRABAJADOR_TELEFONO": (t.telefono or ""),

        "CONTRATO_ID": str(c.id),
        "CONTRATO_FECHA_INGRESO": _fmt_larga(getattr(c, "fecha_inicio", None)),
        "CONTRATO_FECHA_TERMINO": _fmt_larga(getattr(c, "fecha_termino", None)),

        "CARGO_NOMBRE": cargo_nombre,

        "AFP_NOMBRE": afp_nombre,
        "SALUD_NOMBRE": salud_nombre,

        "EMPLEADOR_ID": str(getattr(c, "empleador_id", "") or ""),
        "EMPLEADOR_RAZON_SOCIAL": empleador_razon,
        "EMPLEADOR_RUT": empleador_rut,
        "EMPLEADOR_DIRECCION": empleador_dir,
        "EMPLEADOR_COMUNA": empleador_comuna,
        "REP_LEGAL_NOMBRE": rep_legal_nombre,
        "REP_LEGAL_RUT": rep_legal_rut,

        "OBRA_ID": str(getattr(c, "obra_id", "") or ""),
        "OBRA_NOMBRE": obra_nombre,
        "OBRA_DIRECCION": obra_dir,
        "OBRA_COMUNA": obra_comuna,

        # ✅ CORREGIDO: nombres reales de columnas nuevas en Obra
        "OBRA_ADMINISTRATIVO_NOMBRE": (getattr(obra, "obra_administrativo_nombre", "") or "") if obra else "",
        "OBRA_ADMINISTRATIVO_RUT": (getattr(obra, "obra_administrativo_rut", "") or "") if obra else "",
        "OBRA_ADMINISTRATIVO_CARGO": (getattr(obra, "obra_administrativo_cargo", "") or "") if obra else "",
        "OBRA_ADMINISTRATIVO_TELEFONO": (getattr(obra, "obra_administrativo_telefono", "") or "") if obra else "",
        "OBRA_ADMINISTRATIVO_CORREO": (getattr(obra, "obra_administrativo_correo", "") or "") if obra else "",

        "CAUSAL_RESUMIDA": getattr(causal, "causal_resumida", "") if causal else "",
        "CAUSAL_FINIQUITO": getattr(causal, "causal_finiquito", "") if causal else "",

        "FECHA_TERMINO": _fmt_larga(desv.fecha_termino),
        "FECHA_TERMINO_LARGA": fecha_larga_es(desv.fecha_termino) if getattr(desv, "fecha_termino", None) else "",
        "FECHA_AVISO": _fmt_larga(desv.fecha_aviso),
        "FECHA_AVISO_LARGA": fecha_larga_es(desv.fecha_aviso) if getattr(desv, "fecha_aviso", None) else "",

        "MOTIVO_DETALLE": (desv.motivo_detalle or ""),

        "FINIQUITO_FECHA": _fmt_larga(desv.fecha_termino),
        "FINIQUITO_FECHA_EMISION": _fmt_larga(fecha_emision),
        "FINIQUITO_CAUSAL": getattr(causal, "causal_finiquito", "") if causal else "",

        # defaults (finiquito)
        "FINIQUITO_DIAS_FERIADO_A_PAGO": "",
        "FINIQUITO_FERIADO_PROPORCIONAL": "",
        "FINIQUITO_AVISO_PREVIO": "",
        "FINIQUITO_ANIOS_SERVICIOS": "",
        "FINIQUITO_INDEM_VOLUNTARIA": "",
        "FINIQUITO_MONTO_OTROS": "",
        "FINIQUITO_SUBTOTAL": "",
        "FINIQUITO_APORTE_EMP_SEG_CES": "",
        "FINIQUITO_MONTO_DESCTO_1": "",
        "FINIQUITO_MONTO_DESCTO_2": "",
        "FINIQUITO_DESCTO_1": "",
        "FINIQUITO_DESCTO_2": "",
        "FINIQUITO_TOTAL_DESCTOS": "",
        "FINIQUITO_TOTAL_PAGO": "",

        # causales existentes
        "CAUSAL_7_DETALLE_DIAS": "",
        "CAUSAL_7_FRASE_DIAS": "",
        "CAUSAL_8_FRASE": "",
        "TIEMPO_SERVIDO": "",

        # ✅ NUEVO: bloques carta (para Art. 161)
        "BLOQUE_AVISO_PREVIO": "",
        "BLOQUE_ANIOS_SERVICIO": "",
    }
    return ctx


def _causal_7_context(desv: Desvinculacion) -> dict:
    if not desv.fecha_termino:
        return {}

    month_start, month_end = _month_bounds(desv.fecha_termino)

    rows = (
        db.session.query(Inasistencia.fecha, Inasistencia.motivo)
        .filter(Inasistencia.trabajador_id == desv.trabajador_id)
        .filter(Inasistencia.fecha >= month_start)
        .filter(Inasistencia.fecha <= month_end)
        .filter(
            or_(
                Inasistencia.motivo.ilike("%injustific%"),
                Inasistencia.motivo.ilike("%ausencia_injustificada%"),
                Inasistencia.motivo.ilike("%inasistencia%injustificada%"),
            )
        )
        .order_by(Inasistencia.fecha.asc())
        .all()
    )

    fechas = [r[0] for r in rows if r and r[0]]
    if not fechas:
        return {"CAUSAL_7_DETALLE_DIAS": "", "CAUSAL_7_FRASE_DIAS": ""}

    fechas = sorted(set(fechas))
    runs = _consecutive_runs(fechas)
    run2 = next((r for r in runs if len(r) >= 2), None)
    lunes = [d for d in fechas if d.weekday() == 0]

    if len(fechas) >= 3:
        frase = "tres días en el mismo mes calendario"
        detalle = f"los días {_join_es([str(d.day) for d in fechas])} de {_MESES_ES[fechas[0].month]} de {fechas[0].year}"
        return {"CAUSAL_7_DETALLE_DIAS": detalle, "CAUSAL_7_FRASE_DIAS": frase}

    if run2 and len(fechas) == 2 and (fechas[1] - fechas[0]).days == 1:
        frase = "dos días seguidos en el mismo mes calendario"
        d1, d2 = fechas[0], fechas[1]
        detalle = f"los días {d1.day} y {d2.day} de {_MESES_ES[d1.month]} de {d1.year}"
        return {"CAUSAL_7_DETALLE_DIAS": detalle, "CAUSAL_7_FRASE_DIAS": frase}

    if len(lunes) >= 2 and len(fechas) == 2:
        frase = "dos días lunes en el mismo mes calendario"
        d1, d2 = lunes[0], lunes[1]
        detalle = f"los días lunes {d1.day} y lunes {d2.day} de {_MESES_ES[d1.month]} de {d1.year}"
        return {"CAUSAL_7_DETALLE_DIAS": detalle, "CAUSAL_7_FRASE_DIAS": frase}

    frase = "los días indicados"
    detalle = "los días " + _join_es([str(d.day) for d in fechas]) + f" de {_MESES_ES[fechas[0].month]} de {fechas[0].year}"
    return {"CAUSAL_7_DETALLE_DIAS": detalle, "CAUSAL_7_FRASE_DIAS": frase}


def _causal_8_context(desv: Desvinculacion) -> dict:
    frase = (desv.motivo_detalle or "").strip()
    return {"CAUSAL_8_FRASE": frase}


def _causal_specific_context(desv: Desvinculacion) -> dict:
    causal_id = int(getattr(desv, "causal_id", 0) or 0)
    if causal_id == 7:
        return _causal_7_context(desv)
    if causal_id == 8:
        return _causal_8_context(desv)
    return {}


def build_carta_aviso_docx(desv: Desvinculacion, out_docx: Path, overrides: dict[str, str] | None = None) -> None:
    causal = desv.causal
    if not causal or not getattr(causal, "requiere_carta", True):
        raise RuntimeError("Esta causal no requiere carta aviso.")

    template_name = getattr(causal, "carta_template", None)
    if not template_name:
        raise RuntimeError("Causal requiere carta pero no tiene carta_template.")

    template_path = _resolve_template(template_name)

    # 1) contexto base + causal
    ctx = _base_context(desv) | _causal_specific_context(desv)

    # 2) aplicar overrides (incluye montos desde finiquito si vienen desde UI/DB)
    if overrides:
        ctx = ctx | overrides

    # 3) ✅ inyectar bloques condicionales (necesidades de la empresa)
    ctx = _apply_bloques_carta(ctx)

    _render_docx_template(template_path, _to_str_dict(ctx), out_docx)


def build_finiquito_docx(desv: Desvinculacion, out_docx: Path, overrides: dict[str, str] | None = None) -> None:
    template_name = current_app.config.get("FINIQUITO_TEMPLATE_NAME", "finiquito.docx")
    template_path = _resolve_template(template_name)

    ctx = _base_context(desv) | _causal_specific_context(desv)
    if overrides:
        ctx |= dict(overrides)

    _render_docx_template(template_path, _to_str_dict(ctx), out_docx)


def make_pdf_from_docx(docx_path: Path, pdf_out_dir: Path) -> Path:
    return convert_docx_to_pdf(docx_path, pdf_out_dir)


def upsert_doc_inputs(desvinculacion_id: int, doc_tipo: str, inputs: dict) -> None:
    doc_tipo = (doc_tipo or "").strip().upper()
    if not doc_tipo:
        raise ValueError("doc_tipo vacío")

    row = (
        db.session.query(DesvinculacionDocInputs)
        .filter_by(desvinculacion_id=desvinculacion_id, doc_tipo=doc_tipo)
        .first()
    )

    if row:
        row.inputs = dict(inputs or {})
    else:
        row = DesvinculacionDocInputs(
            desvinculacion_id=desvinculacion_id,
            doc_tipo=doc_tipo,
            inputs=dict(inputs or {}),
        )
        db.session.add(row)


def get_doc_inputs(desvinculacion_id: int, doc_tipo: str) -> dict:
    doc_tipo = (doc_tipo or "").strip().upper()
    row = (
        db.session.query(DesvinculacionDocInputs.inputs)
        .filter_by(desvinculacion_id=desvinculacion_id, doc_tipo=doc_tipo)
        .scalar()
    )
    return row if isinstance(row, dict) else {}