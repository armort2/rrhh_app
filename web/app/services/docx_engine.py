from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# Bookmarks helpers (cargos.docx)
# ============================================================

def _find_bookmark_start(doc: Document, bookmark_name: str):
    for el in doc._element.xpath(".//*[local-name()='bookmarkStart']"):
        if el.get(qn("w:name")) == bookmark_name:
            return el
    return None


def _find_bookmark_end_by_id(doc: Document, bookmark_id: str):
    for el in doc._element.xpath(".//*[local-name()='bookmarkEnd']"):
        if el.get(qn("w:id")) == bookmark_id:
            return el
    return None


# ============================================================
# Placeholders: reemplazo robusto sin “negritar todo”
# (mantiene formato del run donde está el placeholder)
# ============================================================

def _iter_all_story_parts(doc: Document) -> Iterable[Tuple[str, object]]:
    yield ("body", doc)
    for i, sec in enumerate(doc.sections):
        yield (f"section[{i}].header", sec.header)
        yield (f"section[{i}].footer", sec.footer)

        if getattr(sec, "first_page_header", None) is not None:
            yield (f"section[{i}].first_page_header", sec.first_page_header)
        if getattr(sec, "first_page_footer", None) is not None:
            yield (f"section[{i}].first_page_footer", sec.first_page_footer)
        if getattr(sec, "even_page_header", None) is not None:
            yield (f"section[{i}].even_page_header", sec.even_page_header)
        if getattr(sec, "even_page_footer", None) is not None:
            yield (f"section[{i}].even_page_footer", sec.even_page_footer)


def _replace_in_paragraph_runs(paragraph, mapping: Dict[str, str]) -> bool:
    """
    Estrategia:
    1) Reemplazo directo dentro de cada run (ideal: mantiene negrita solo del placeholder).
    2) Fallback: si el placeholder está partido entre runs, merge CONTROLADO solo de los runs
       involucrados en la ocurrencia (sin colapsar todo el párrafo).
    """
    if not paragraph.runs:
        return False

    changed_any = False

    # 1) Reemplazo directo por run
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = run.text
        for k, v in mapping.items():
            if k in new_text:
                new_text = new_text.replace(k, v)
        if new_text != run.text:
            run.text = new_text
            changed_any = True

    # 2) Fallback: placeholders partidos entre runs
    runs = paragraph.runs
    full = "".join(r.text or "" for r in runs)
    if not full:
        return changed_any

    for token, value in mapping.items():
        runs = paragraph.runs
        full = "".join(r.text or "" for r in runs)

        while token in full:
            start_idx = full.find(token)
            end_idx = start_idx + len(token)

            pos = 0
            start_run_i = None
            end_run_i = None
            start_off = 0
            end_off = 0

            for i, r in enumerate(runs):
                t = r.text or ""
                next_pos = pos + len(t)

                if start_run_i is None and start_idx < next_pos:
                    start_run_i = i
                    start_off = start_idx - pos

                if start_run_i is not None and end_idx <= next_pos:
                    end_run_i = i
                    end_off = end_idx - pos
                    break

                pos = next_pos

            if start_run_i is None or end_run_i is None:
                break

            before = (runs[start_run_i].text or "")[:start_off]
            after = (runs[end_run_i].text or "")[end_off:]

            runs[start_run_i].text = before + value + after
            for j in range(start_run_i + 1, end_run_i + 1):
                runs[j].text = ""

            changed_any = True
            full = "".join(r.text or "" for r in paragraph.runs)

    return changed_any


def _replace_in_table(table, mapping: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph_runs(p, mapping)
            for t in cell.tables:
                _replace_in_table(t, mapping)


def _replace_in_textboxes_via_xml(doc: Document, mapping: Dict[str, str]) -> None:
    texts = doc._element.xpath(".//*[local-name()='txbxContent']//*[local-name()='t']")
    if not texts:
        return

    for t in texts:
        txt = t.text or ""
        if not txt:
            continue
        new = txt
        for k, v in mapping.items():
            if k in new:
                new = new.replace(k, v)
        if new != txt:
            t.text = new


def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
    for _, part in _iter_all_story_parts(doc):
        for p in getattr(part, "paragraphs", []):
            _replace_in_paragraph_runs(p, mapping)

        for table in getattr(part, "tables", []):
            _replace_in_table(table, mapping)

    _replace_in_textboxes_via_xml(doc, mapping)


# ============================================================
# Opción C: Merge de numeración (listas) entre docs
# ============================================================

def _get_numbering_root(doc: Document):
    part = doc.part.numbering_part
    return part.element


def _get_attr_int_local(el, local_attr_name: str) -> Optional[int]:
    if el is None:
        return None
    for k, v in el.attrib.items():
        if k.endswith("}" + local_attr_name) or k.endswith(":" + local_attr_name) or k == local_attr_name:
            try:
                return int(v)
            except Exception:
                return None
    return None


def _existing_num_ids(num_root) -> List[int]:
    ids: List[int] = []
    for num in num_root.xpath(".//*[local-name()='num']"):
        v = _get_attr_int_local(num, "numId")
        if v is not None:
            ids.append(v)
    return ids


def _existing_abstract_ids(num_root) -> List[int]:
    ids: List[int] = []
    for a in num_root.xpath(".//*[local-name()='abstractNum']"):
        v = _get_attr_int_local(a, "abstractNumId")
        if v is not None:
            ids.append(v)
    return ids


def _next_free_id(used: List[int], start: int = 1000) -> int:
    s = set(used)
    i = start
    while i in s:
        i += 1
    return i


def _paragraph_num_id(p_xml) -> Optional[int]:
    numId_nodes = p_xml.xpath(".//*[local-name()='numPr']/*[local-name()='numId']")
    if not numId_nodes:
        return None

    v = None
    for k, vv in numId_nodes[0].attrib.items():
        if k.endswith("}val") or k.endswith(":val") or k == "val":
            v = vv
            break
    if v is None:
        return None
    try:
        return int(v)
    except Exception:
        return None


def _set_paragraph_num_id(p_xml, new_num_id: int) -> None:
    numId_nodes = p_xml.xpath(".//*[local-name()='numPr']/*[local-name()='numId']")
    if not numId_nodes:
        return
    numId_nodes[0].set(qn("w:val"), str(new_num_id))


def _find_num_def(num_root, num_id: int):
    nodes = num_root.xpath(
        f".//*[local-name()='num' and @*[local-name()='numId']='{num_id}']"
    )
    return nodes[0] if nodes else None


def _find_abstract_def(num_root, abstract_id: int):
    nodes = num_root.xpath(
        f".//*[local-name()='abstractNum' and @*[local-name()='abstractNumId']='{abstract_id}']"
    )
    return nodes[0] if nodes else None


def _num_abstract_id(num_xml) -> Optional[int]:
    abs_nodes = num_xml.xpath(".//*[local-name()='abstractNumId']")
    if not abs_nodes:
        return None

    v = abs_nodes[0].get(qn("w:val"))
    if v is None:
        for k, vv in abs_nodes[0].attrib.items():
            if k.endswith("}val") or k.endswith(":val") or k == "val":
                v = vv
                break
    if v is None:
        return None

    try:
        return int(v)
    except Exception:
        return None


def _set_num_abstract_id(num_xml, new_abs_id: int) -> None:
    abs_nodes = num_xml.xpath(".//*[local-name()='abstractNumId']")
    if not abs_nodes:
        return
    abs_nodes[0].set(qn("w:val"), str(new_abs_id))


def merge_numbering_for_paragraphs(
    dst_doc: Document,
    src_doc: Document,
    paragraph_xml_list: List[object],
) -> None:
    dst_num = _get_numbering_root(dst_doc)
    src_num = _get_numbering_root(src_doc)

    needed_num_ids: List[int] = []
    for p in paragraph_xml_list:
        nid = _paragraph_num_id(p)
        if nid is not None:
            needed_num_ids.append(nid)

    needed_num_ids = sorted(set(needed_num_ids))
    if not needed_num_ids:
        return

    used_dst_num = _existing_num_ids(dst_num)
    used_dst_abs = _existing_abstract_ids(dst_num)

    num_map: Dict[int, int] = {}
    abs_map: Dict[int, int] = {}

    for old_num_id in needed_num_ids:
        src_num_def = _find_num_def(src_num, old_num_id)
        if src_num_def is None:
            continue

        old_abs_id = _num_abstract_id(src_num_def)

        new_num_id = _next_free_id(used_dst_num, start=1000)
        used_dst_num.append(new_num_id)
        num_map[old_num_id] = new_num_id

        new_num_def = deepcopy(src_num_def)
        new_num_def.set(qn("w:numId"), str(new_num_id))

        if old_abs_id is not None:
            if old_abs_id in abs_map:
                new_abs_id = abs_map[old_abs_id]
            else:
                new_abs_id = _next_free_id(used_dst_abs, start=1000)
                used_dst_abs.append(new_abs_id)
                abs_map[old_abs_id] = new_abs_id

                src_abs_def = _find_abstract_def(src_num, old_abs_id)
                if src_abs_def is not None:
                    new_abs_def = deepcopy(src_abs_def)
                    new_abs_def.set(qn("w:abstractNumId"), str(new_abs_id))
                    dst_num.append(new_abs_def)

            _set_num_abstract_id(new_num_def, new_abs_id)

        dst_num.append(new_num_def)

    for p in paragraph_xml_list:
        old = _paragraph_num_id(p)
        if old is not None and old in num_map:
            _set_paragraph_num_id(p, num_map[old])


# ============================================================
# Inserción del bloque de cargo (bookmark) + merge numbering
# ============================================================

def _collect_bookmark_paragraphs_xml(source_doc: Document, bookmark_name: str) -> List[object]:
    start = _find_bookmark_start(source_doc, bookmark_name)
    if start is None:
        raise ValueError(f"No se encontró el marcador '{bookmark_name}' en cargos.docx")

    bookmark_id = start.get(qn("w:id"))
    if not bookmark_id:
        raise ValueError(f"Marcador '{bookmark_name}' inválido: sin w:id")

    end = _find_bookmark_end_by_id(source_doc, bookmark_id)
    if end is None:
        raise ValueError(f"No se encontró el bookmarkEnd para '{bookmark_name}' (id={bookmark_id})")

    all_els = list(source_doc._element.iter())
    try:
        i_start = all_els.index(start)
        i_end = all_els.index(end)
    except ValueError as e:
        raise ValueError(f"No se pudo localizar el rango del marcador '{bookmark_name}'.") from e

    if i_end < i_start:
        i_start, i_end = i_end, i_start

    paras = []
    seen = set()

    for el in all_els[i_start:i_end + 1]:
        p = el
        while p is not None and p.tag != qn("w:p"):
            p = p.getparent()
        if p is not None and p.tag == qn("w:p"):
            pid = id(p)
            if pid not in seen:
                seen.add(pid)
                paras.append(p)

    cleaned = []
    for p in paras:
        txt = "".join(t.text or "" for t in p.xpath(".//*[local-name()='t']"))
        if txt.strip():
            cleaned.append(p)

    if not cleaned:
        raise ValueError(f"El marcador '{bookmark_name}' existe pero no contiene párrafos insertables.")

    return cleaned


def insert_text_block_at_placeholder(
    target_doc: Document,
    placeholder: str,
    source_doc: Document,
    bookmark_name: str,
) -> None:
    placeholder_p = None
    for p in target_doc.paragraphs:
        if placeholder in p.text:
            placeholder_p = p
            break
    if placeholder_p is None:
        raise ValueError(f"No se encontró el placeholder '{placeholder}' en contrato_base.docx")

    src_paras_xml = _collect_bookmark_paragraphs_xml(source_doc, bookmark_name)
    copied_paras = [deepcopy(p) for p in src_paras_xml]

    merge_numbering_for_paragraphs(
        dst_doc=target_doc,
        src_doc=source_doc,
        paragraph_xml_list=copied_paras,
    )

    parent = placeholder_p._element.getparent()
    insert_at = parent.index(placeholder_p._element)

    for p_xml in copied_paras:
        parent.insert(insert_at, p_xml)
        insert_at += 1

    parent.remove(placeholder_p._element)


# ============================================================
# POST-PROCESO: forzar viñetas en artículos Séptimo y Octavo
# ============================================================

def _normalize_heading(s: str) -> str:
    """
    Normaliza texto para comparar encabezados (sin tildes, sin puntos, upper).
    """
    if not s:
        return ""
    t = s.strip().upper()
    # normalización mínima (sin depender de libs externas)
    t = (
        t.replace("Á", "A").replace("É", "E").replace("Í", "I")
         .replace("Ó", "O").replace("Ú", "U").replace("Ü", "U")
         .replace("Ñ", "N")
    )
    return t


def _find_article_index(doc: Document, article_token: str) -> Optional[int]:
    target = _normalize_heading(article_token)
    for i, p in enumerate(doc.paragraphs):
        txt = _normalize_heading(p.text)
        if txt.startswith(target):
            return i
    return None


def _has_numpr(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='numPr']"))  # pylint: disable=protected-access


def _ensure_bullet_numbering(doc: Document) -> int:
    """
    Crea (si es necesario) una definición propia de viñeta (bullet) y retorna su numId.
    Se crea con IDs altos para evitar colisiones.

    NOTA: NO forzamos Wingdings, porque puede renderizar "•" como otros glifos (ej: ❿).
    Dejamos que Word elija la fuente correcta para bullets.
    """
    num_root = _get_numbering_root(doc)
    used_num = _existing_num_ids(num_root)
    used_abs = _existing_abstract_ids(num_root)

    bullet_abs_id = _next_free_id(used_abs, start=5000)
    bullet_num_id = _next_free_id(used_num, start=5000)

    # abstractNum
    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(qn("w:abstractNumId"), str(bullet_abs_id))

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")

    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")

    # Bullet estándar Unicode (dejando que Word use la fuente apropiada)
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "•")

    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")

    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    lvl.append(start)
    lvl.append(numFmt)
    lvl.append(lvlText)
    lvl.append(lvlJc)
    lvl.append(pPr)

    # Importante: NO agregamos w:rPr con Wingdings/Symbol para no “mutar” el bullet.
    abstractNum.append(lvl)
    num_root.append(abstractNum)

    # num
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(bullet_num_id))
    abstractNumId = OxmlElement("w:abstractNumId")
    abstractNumId.set(qn("w:val"), str(bullet_abs_id))
    num.append(abstractNumId)

    num_root.append(num)

    return bullet_num_id



def _set_paragraph_to_bullet(paragraph, bullet_num_id: int) -> None:
    """
    Fuerza el párrafo a usar nuestra numeración tipo bullet.
    Mantiene numPr existente, pero ajusta numId e ilvl=0.
    """
    numPr_nodes = paragraph._p.xpath(".//*[local-name()='numPr']")  # pylint: disable=protected-access
    if not numPr_nodes:
        return
    numPr = numPr_nodes[0]

    # asegurar ilvl
    ilvl_nodes = numPr.xpath("./*[local-name()='ilvl']")
    if ilvl_nodes:
        ilvl_nodes[0].set(qn("w:val"), "0")
    else:
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numPr.insert(0, ilvl)

    # set numId
    numId_nodes = numPr.xpath("./*[local-name()='numId']")
    if numId_nodes:
        numId_nodes[0].set(qn("w:val"), str(bullet_num_id))
    else:
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(bullet_num_id))
        numPr.append(numId)


def post_process_fix_bullets_in_articles(doc: Document) -> None:
    """
    Repara las listas de los artículos Séptimo y Octavo:
    - detecta rango del artículo (desde encabezado hasta el siguiente encabezado)
    - solo afecta párrafos con numPr (listas)
    - los fuerza a una definición bullet propia (numId nuevo)
    """
    bullet_num_id = _ensure_bullet_numbering(doc)

    # Define rangos que queremos reparar
    targets = [
        ("SEPTIMO", "OCTAVO"),
        ("SÉPTIMO", "OCTAVO"),
        ("OCTAVO", "NOVENO"),
        ("OCTAVO", "DÉCIMO"),
        ("OCTAVO", "DECIMO"),
    ]

    # Encontrar índice real de "Séptimo" y "Octavo"
    i_sept = _find_article_index(doc, "SÉPTIMO")
    if i_sept is None:
        i_sept = _find_article_index(doc, "SEPTIMO")

    i_oct = _find_article_index(doc, "OCTAVO")

    # Si no encontramos, no tocamos nada
    if i_sept is None and i_oct is None:
        return

    def _find_next_heading(from_idx: int, candidates: List[str]) -> Optional[int]:
        for j in range(from_idx + 1, len(doc.paragraphs)):
            t = _normalize_heading(doc.paragraphs[j].text)
            for c in candidates:
                if t.startswith(_normalize_heading(c)):
                    return j
        return None

    # Reparar Séptimo (hasta Octavo)
    if i_sept is not None:
        end = i_oct if i_oct is not None else _find_next_heading(i_sept, ["OCTAVO", "NOVENO", "DÉCIMO", "DECIMO"])
        if end is None:
            end = len(doc.paragraphs)

        for p in doc.paragraphs[i_sept:end]:
            if _has_numpr(p):
                _set_paragraph_to_bullet(p, bullet_num_id)

    # Reparar Octavo (hasta Noveno/Décimo)
    if i_oct is not None:
        end = _find_next_heading(i_oct, ["NOVENO", "DÉCIMO", "DECIMO", "UNDÉCIMO", "UNDECIMO"])
        if end is None:
            end = len(doc.paragraphs)

        for p in doc.paragraphs[i_oct:end]:
            if _has_numpr(p):
                _set_paragraph_to_bullet(p, bullet_num_id)


# ============================================================
# Pipeline final
# ============================================================

def build_contract_docx(
    base_template_path: Path,
    cargos_library_path: Path,
    cargo_id: int,
    variables: Dict[str, str],
    output_path: Path,
    cargo_bookmark_template: str = "CARGO_{id}__DESCRIPCION",
    cargo_placeholder: str = "{{DESCRIPCION_CARGO}}",
) -> Path:
    base_doc = Document(str(base_template_path))
    cargos_doc = Document(str(cargos_library_path))

    bookmark_name = cargo_bookmark_template.format(id=cargo_id)

    replace_placeholders(base_doc, variables)

    insert_text_block_at_placeholder(
        target_doc=base_doc,
        placeholder=cargo_placeholder,
        source_doc=cargos_doc,
        bookmark_name=bookmark_name,
    )

    # POST-PROCESO: reparar viñetas en Séptimo y Octavo
    post_process_fix_bullets_in_articles(base_doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    base_doc.save(str(output_path))
    return output_path
